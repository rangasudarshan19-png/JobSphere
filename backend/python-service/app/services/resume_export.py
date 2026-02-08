"""
Resume Export Service - Generate PDF and DOCX files
Supports multiple templates and formats
"""

from typing import Dict, Optional
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64


class ResumeExporter:
    """Generate resume files in PDF and DOCX formats"""
    
    TEMPLATES = {
        "Professional": {
            "font": "Calibri",
            "heading_size": 16,
            "body_size": 11,
            "color": (47, 84, 150),  # Professional blue
            "spacing": 1.15
        },
        "Modern": {
            "font": "Arial",
            "heading_size": 18,
            "body_size": 11,
            "color": (99, 102, 241),  # Modern purple
            "spacing": 1.2
        },
        "Creative": {
            "font": "Georgia",
            "heading_size": 17,
            "body_size": 11,
            "color": (236, 72, 153),  # Creative pink
            "spacing": 1.3
        },
        "ATS-Optimized": {
            "font": "Times New Roman",
            "heading_size": 14,
            "body_size": 11,
            "color": (0, 0, 0),  # Black for ATS
            "spacing": 1.0
        },
        "Executive": {
            "font": "Garamond",
            "heading_size": 16,
            "body_size": 11,
            "color": (31, 41, 55),  # Executive dark gray
            "spacing": 1.15
        }
    }
    
    async def generate_docx(
        self,
        resume_data: Dict,
        template_style: str = "Professional"
    ) -> bytes:
        """
        Generate DOCX resume file
        
        Args:
            resume_data: Resume content dictionary
            template_style: Template name (Professional, Modern, etc.)
            
        Returns:
            bytes: DOCX file content
        """
        try:
            # Get template config
            config = self.TEMPLATES.get(template_style, self.TEMPLATES["Professional"])
            
            # Create document
            doc = Document()
            
            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = config["font"]
            font.size = Pt(config["body_size"])
            
            # Extract data
            contact = resume_data.get("contact", {})
            summary = resume_data.get("summary", "")
            experience = resume_data.get("experience", [])
            education = resume_data.get("education", [])
            skills = resume_data.get("skills", [])
            projects = resume_data.get("projects", [])
            certifications = resume_data.get("certifications", [])
            
            # Add profile picture if provided
            profile_picture = contact.get("profile_picture")
            if profile_picture:
                try:
                    # If base64 encoded
                    if profile_picture.startswith('data:image'):
                        # Extract base64 data
                        image_data = profile_picture.split(',')[1]
                        image_bytes = base64.b64decode(image_data)
                        image_stream = BytesIO(image_bytes)
                        doc.add_picture(image_stream, width=Inches(1.5))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    print(f"[SYMBOL]️ Could not add profile picture: {e}")
            
            # Header - Name
            name = doc.add_heading(contact.get("full_name", ""), 0)
            name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in name.runs:
                run.font.color.rgb = RGBColor(*config["color"])
                run.font.size = Pt(config["heading_size"] + 4)
            
            # Contact Info
            contact_info = []
            if contact.get("email"):
                contact_info.append(contact["email"])
            if contact.get("phone"):
                contact_info.append(contact["phone"])
            if contact.get("location"):
                contact_info.append(contact["location"])
            if contact.get("linkedin"):
                contact_info.append(contact["linkedin"])
            
            if contact_info:
                p = doc.add_paragraph(" | ".join(contact_info))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
            
            doc.add_paragraph()  # Spacing
            
            # Professional Summary
            if summary:
                self._add_section(doc, "Professional Summary", summary, config)
            
            # Experience
            if experience:
                heading = doc.add_heading("Work Experience", level=1)
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(*config["color"])
                
                for exp in experience:
                    # Job title and company
                    job_para = doc.add_paragraph()
                    job_run = job_para.add_run(f"{exp.get('title', '')} - {exp.get('company', '')}")
                    job_run.bold = True
                    job_run.font.size = Pt(12)
                    
                    # Duration
                    duration = f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}"
                    date_para = doc.add_paragraph(duration)
                    for run in date_para.runs:
                        run.italic = True
                        run.font.size = Pt(10)
                    
                    # Description
                    if exp.get('description'):
                        if isinstance(exp['description'], list):
                            for item in exp['description']:
                                p = doc.add_paragraph(item, style='List Bullet')
                                p.paragraph_format.left_indent = Inches(0.25)
                        else:
                            doc.add_paragraph(exp['description'])
                    
                    doc.add_paragraph()  # Spacing between jobs
            
            # Education
            if education:
                heading = doc.add_heading("Education", level=1)
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(*config["color"])
                
                for edu in education:
                    degree_para = doc.add_paragraph()
                    degree_run = degree_para.add_run(f"{edu.get('degree', '')} - {edu.get('institution', '')}")
                    degree_run.bold = True
                    degree_run.font.size = Pt(12)
                    
                    if edu.get('graduation_date'):
                        date_para = doc.add_paragraph(edu['graduation_date'])
                        for run in date_para.runs:
                            run.italic = True
                            run.font.size = Pt(10)
                    
                    doc.add_paragraph()  # Spacing
            
            # Skills
            if skills:
                heading = doc.add_heading("Skills", level=1)
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(*config["color"])
                
                if isinstance(skills, list):
                    skills_text = ", ".join(skills)
                else:
                    skills_text = skills
                doc.add_paragraph(skills_text)
                doc.add_paragraph()
            
            # Projects
            if projects:
                heading = doc.add_heading("Projects", level=1)
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(*config["color"])
                
                for project in projects:
                    project_para = doc.add_paragraph()
                    project_run = project_para.add_run(project.get('name', ''))
                    project_run.bold = True
                    project_run.font.size = Pt(12)
                    
                    if project.get('description'):
                        doc.add_paragraph(project['description'])
                    
                    doc.add_paragraph()  # Spacing
            
            # Certifications
            if certifications:
                heading = doc.add_heading("Certifications", level=1)
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(*config["color"])
                
                if isinstance(certifications, list):
                    for cert in certifications:
                        doc.add_paragraph(cert, style='List Bullet')
                else:
                    doc.add_paragraph(certifications)
            
            # Save to BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return buffer.getvalue()
            
        except Exception as e:
            print(f"[SYMBOL] DOCX generation error: {e}")
            raise e
    
    async def generate_pdf(
        self,
        resume_data: Dict,
        template_style: str = "Professional"
    ) -> bytes:
        """
        Generate PDF resume file
        
        Args:
            resume_data: Resume content dictionary
            template_style: Template name
            
        Returns:
            bytes: PDF file content
        """
        try:
            print("[INFO] Starting PDF generation...")
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
                from reportlab.lib.enums import TA_CENTER, TA_LEFT
                print("[INFO] Reportlab imported successfully")
            except ImportError as ie:
                print(f"[ERROR] Failed to import reportlab: {ie}")
                raise Exception(f"PDF library not available: {ie}")
            
            # Get template config
            config = self.TEMPLATES.get(template_style, self.TEMPLATES["Professional"])
            
            # Create PDF buffer
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter,
                                   topMargin=0.5*inch, bottomMargin=0.5*inch,
                                   leftMargin=0.75*inch, rightMargin=0.75*inch)
            
            # Styles
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=config["heading_size"] + 4,
                textColor=config["color"],
                alignment=TA_CENTER,
                spaceAfter=6
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=config["heading_size"],
                textColor=config["color"],
                spaceAfter=6,
                spaceBefore=12
            )
            
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['Normal'],
                fontSize=config["body_size"],
                spaceAfter=6
            )
            
            contact_style = ParagraphStyle(
                'Contact',
                parent=styles['Normal'],
                fontSize=10,
                alignment=TA_CENTER,
                spaceAfter=12
            )
            
            # Build content
            story = []
            
            # Extract data - handle both dict and object formats
            contact = resume_data.get("contact", {})
            summary = resume_data.get("summary", "") or resume_data.get("professional_summary", "")
            experience = resume_data.get("experience", []) or resume_data.get("work_experience", [])
            education = resume_data.get("education", [])
            skills = resume_data.get("skills", [])
            projects = resume_data.get("projects", [])
            certifications = resume_data.get("certifications", [])
            
            # Name
            full_name = contact.get("full_name", "") if isinstance(contact, dict) else str(contact)
            story.append(Paragraph(str(full_name), title_style))
            
            # Contact info
            contact_info = []
            if isinstance(contact, dict):
                if contact.get("email"):
                    contact_info.append(str(contact["email"]))
                if contact.get("phone"):
                    contact_info.append(str(contact["phone"]))
                if contact.get("location"):
                    contact_info.append(str(contact["location"]))
                if contact.get("linkedin"):
                    contact_info.append(str(contact["linkedin"]))
            
            if contact_info:
                story.append(Paragraph(" | ".join(contact_info), contact_style))
            
            story.append(Spacer(1, 0.2*inch))
            
            # Summary
            if summary:
                story.append(Paragraph("Professional Summary", heading_style))
                story.append(Paragraph(str(summary), body_style))
                story.append(Spacer(1, 0.1*inch))
            
            # Experience
            if experience:
                story.append(Paragraph("Work Experience", heading_style))
                if isinstance(experience, list):
                    for exp in experience:
                        if isinstance(exp, dict):
                            title_company = f"<b>{exp.get('title', '') or exp.get('job_title', '')}</b> - {exp.get('company', '')}"
                            story.append(Paragraph(title_company, body_style))
                            
                            duration = f"<i>{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}</i>"
                            story.append(Paragraph(duration, body_style))
                            
                            if exp.get('description'):
                                if isinstance(exp['description'], list):
                                    for item in exp['description']:
                                        story.append(Paragraph(f"• {item}", body_style))
                                else:
                                    story.append(Paragraph(str(exp['description']), body_style))
                            
                            story.append(Spacer(1, 0.1*inch))
            
            # Education
            if education:
                story.append(Paragraph("Education", heading_style))
                if isinstance(education, list):
                    for edu in education:
                        if isinstance(edu, dict):
                            degree_school = f"<b>{edu.get('degree', '')}</b> - {edu.get('institution', '')}"
                            story.append(Paragraph(degree_school, body_style))
                            
                            if edu.get('graduation_date') or edu.get('year'):
                                story.append(Paragraph(f"<i>{edu.get('graduation_date', edu.get('year', ''))}</i>", body_style))
                            
                            story.append(Spacer(1, 0.1*inch))
            
            # Skills
            if skills:
                story.append(Paragraph("Skills", heading_style))
                if isinstance(skills, list):
                    skills_text = ", ".join(str(s) for s in skills)
                else:
                    skills_text = str(skills)
                story.append(Paragraph(skills_text, body_style))
                story.append(Spacer(1, 0.1*inch))
            
            # Projects
            if projects:
                story.append(Paragraph("Projects", heading_style))
                if isinstance(projects, list):
                    for project in projects:
                        if isinstance(project, dict):
                            story.append(Paragraph(f"<b>{project.get('name', '')}</b>", body_style))
                            if project.get('description'):
                                story.append(Paragraph(str(project['description']), body_style))
                            story.append(Spacer(1, 0.1*inch))
            
            # Certifications
            if certifications:
                story.append(Paragraph("Certifications", heading_style))
                if isinstance(certifications, list):
                    for cert in certifications:
                        story.append(Paragraph(f"• {cert}", body_style))
                else:
                    story.append(Paragraph(str(certifications), body_style))
            
            # Build PDF
            doc.build(story)
            buffer.seek(0)
            
            print("[SYMBOL] PDF generated successfully")
            return buffer.getvalue()
            
        except ImportError:
            # If reportlab not installed, return DOCX as fallback
            print("[SYMBOL]️ reportlab not installed. Falling back to DOCX format.")
            return await self.generate_docx(resume_data, template_style)
        except Exception as e:
            print(f"[SYMBOL] PDF generation error: {str(e)[:200]}")
            # Fallback to DOCX
            print("   Falling back to DOCX format...")
            try:
                return await self.generate_docx(resume_data, template_style)
            except Exception as docx_error:
                print(f"[SYMBOL] DOCX fallback also failed: {str(docx_error)[:200]}")
                raise Exception(f"Failed to generate resume: {str(e)}")
    
    def _add_section(self, doc, title: str, content: str, config: Dict):
        """Helper to add a section to document"""
        heading = doc.add_heading(title, level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*config["color"])
        
        doc.add_paragraph(content)
        doc.add_paragraph()  # Spacing


# Global instance
resume_exporter = ResumeExporter()
